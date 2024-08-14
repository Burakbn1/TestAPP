from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docxtpl import DocxTemplate, InlineImage
import os, sys



class degisken_atama:
    def __init__(self, dosya_adi, resim_sayisi):
        super().__init__()
        self.dosya_adi = dosya_adi
        self.resim_sayisi = resim_sayisi
    def degisken_konum_atama(self, doc, insert_after,resim_sayisi):
        # Yeni sayfaya geçiş yapmak için önceki paragrafın sonuna sayfa sonu ekle
        

        # Tabloya veri ekle RESİM SAYİSİ GELMELİ
        for islem_sayisi in range(resim_sayisi, 0,-1): 
            
            run = insert_after.add_run()
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            # Yeni sayfada belgeye bir başlık ekle ve ortala
            heading = doc.add_paragraph()
            
            heading_run = heading.add_run('Drawing No. <drawno1> Rev.: <drawrev1> Page.: 1-')
            heading_run.font.name = 'Arial'
            heading_run.font.size = Pt(14)
            
            # Boş bir paragraf ekle
            doc.add_paragraph(' ')

            # Belgeye bir tablo ekle (1 satır, 1 sütun)
            table = doc.add_table(rows=1, cols=1)

            # Tablo hücresini boyutlandır
            table.cell(0, 0).width = Mm(180)
            table.cell(0, 0).height = Mm(210)
            cell = table.cell(0, 0)
            cell.text = '/resim'+ str(islem_sayisi)
            # Metni ortala
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0]
            run.font.size = Pt(12)

            # Tabloya kenarlık ekle
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.append(tblPr)

            tblBorders = OxmlElement('w:tblBorders')

            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tblBorders.append(border)

            tblPr.append(tblBorders)
            
            # Yeni içeriği doğru konuma ekle
            insert_after._element.addnext(heading._element)
            
            heading._element.addnext(table._element)
            
            

            

        doc.save('zson_hal.docx')
        self.doc= Document('zson_hal.docx')


    def resim_url_alma(self, resim_url):
        
        gelen_url= resim_url
        #print("aranan url="+resim_url[2])
        return gelen_url

    def islem(self, gelen_url, olusturulan_dosya_adi):
            #doc.save('zson_hal.docx')
        docx= DocxTemplate(olusturulan_dosya_adi)
        for i in range(self.resim_sayisi):
            docx= DocxTemplate(olusturulan_dosya_adi)
            sorgulanan_dosya_adi= 'resim'+ str(i+1)
            context = {
                sorgulanan_dosya_adi: InlineImage(docx, f'{gelen_url[i]}', width=Mm(180), height=Mm(200))
            }
            print(gelen_url[i])
            print(f"Adding image {i+1}")
            print("Adding image"+ str(i+1))

            docx.render(context)
            docx.save(olusturulan_dosya_adi)
            print("dosyakaydıBaşarılı")

    def degisken_deger_atama():
        os.chdir(sys.path[0])
        doc = DocxTemplate('z2.hal.docx')

        place_holder1 = InlineImage(doc, 'resim1.png', width=Mm(180), height=Mm(200))

        context = {
            'resim1': place_holder1
        }

        doc.render(context)
        #doc.save('zson_hal.docx')
        print("deneme başarılı")
