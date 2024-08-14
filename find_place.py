from docx import Document
import new_page
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Pt, Mm

# Mevcut bir belgeyi aç


# Değiştirmek istediğiniz metin

def find_place_pics(dosya_adi, resim_sayisi, resim_url):
   
    print(dosya_adi+"dosya adi")
    olusturulan_dosya_adi= dosya_adi+ '.docx'
    
    doc = Document(olusturulan_dosya_adi)
    
    old_text = 'form_noktasi'
    # Belgedeki tüm paragrafları kontrol et
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            # Metni bul ve sil
            paragraph.text = paragraph.text.replace(old_text, '')
            
            # Yeni içeriği ekle
            degisken_atama_instance = new_page.degisken_atama(dosya_adi, resim_sayisi)
            degisken_atama_instance.degisken_konum_atama(doc, paragraph, resim_sayisi)
            doc.save(olusturulan_dosya_adi)
            degisken_atama_instance.islem(resim_url, olusturulan_dosya_adi)
            print("find_place_pics, break üstü")
            break


def create_place_pic_spot(dosya_adi, resim_sayisi, resim_url):
    doc= Document(dosya_adi+ '.docx')
    i=1
    print(dosya_adi)
    for i in range(resim_sayisi):
        old_text = '/resim'+str(i+1)
        print(old_text)
        # Belgedeki tüm paragrafları kontrol et
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if old_text in cell.text:
                     # Metni bul ve sil
                        cell.text = ""

                        # Resmi ekle
                        run = cell.paragraphs[0].add_run()
                        run.add_picture(f'{resim_url[i]}', width=Mm(180), height=Mm(200))
                        doc.save('create_place_pic_spot.docx')
                        #paragraph.text = paragraph.text.replace(old_text, '')
                        #new_page.degisken_atama.islem(resim_url, dosya_adi)
                    
                        print("create_place_pic_spot")
                        break

    # Belgeyi kaydet
    #doc.save('create_place_pic_spot.docx')

def find_place_toc():
    doc= Document('aTemprise.docx')
    old_text= 'Table of Contents'

    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            #table_of_contents.degisken_deger_atama()
            

            break
        
    
